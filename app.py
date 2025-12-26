from flask import Flask, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

app = Flask(__name__)

@app.route("/")
def generate_doc():

    # Create document
    doc = Document()

    # Helper function for centered headings
    def add_centered(text, size=12, bold=True):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Title Section
    add_centered("FORM ‘A’", 12)
    add_centered("MEDIATION APPLICATION FORM", 12)
    add_centered("[REFER RULE 3(1)]", 11)
    add_centered("Mumbai District Legal Services Authority", 11)
    add_centered("City Civil Court, Mumbai", 11)

    doc.add_paragraph("")

    # Main Table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header Row
    hdr = table.rows[0].cells
    hdr[0].merge(hdr[2])
    hdr[0].text = "DETAILS OF PARTIES:"
    hdr[0].paragraphs[0].runs[0].bold = True

    # Row helper
    def add_row(c1="", c2="", c3=""):
        row = table.add_row().cells
        row[0].text = c1
        row[1].text = c2
        row[2].text = c3
        for cell in row:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Reduce spacing in first column
        row[0].width = Inches(0.6)
        row[1].width = Inches(2.5)
        row[2].width = Inches(2.5)

    # Applicant Details
    add_row("1", "Name of Applicant", "")
    add_row("", "Address and contact details of Applicant", "")
    add_row("", "REGISTERED ADDRESS:", "")
    add_row("", "CORRESPONDENCE BRANCH ADDRESS:", "")
    add_row("", "Telephone No.", "")
    add_row("", "Mobile No.", "")
    add_row("", "Email ID", "")

    # Opposite Party
    add_row("2", "Name, Address and Contact details of Opposite Party:", "")
    add_row("", "Name", "")
    add_row("", "REGISTERED ADDRESS:", "")
    add_row("", "CORRESPONDENCE ADDRESS:", "")
    add_row("", "Telephone No.", "")
    add_row("", "Mobile No.", "")
    add_row("", "Email ID", "")

    # Dispute Section
    row = table.add_row().cells
    row[0].merge(row[2])
    row[0].text = "DETAILS OF DISPUTE:"
    row[0].paragraphs[0].runs[0].bold = True

    row = table.add_row().cells
    row[0].merge(row[2])
    p = row[0].paragraphs[0]
    run = p.add_run("THE COMM. COURTS (PRE-INSTITUTION) SETTLEMENT RULES, 2018")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = table.add_row().cells
    row[0].merge(row[2])
    row[0].text = "Nature of disputes as per section 2(1)(c) of the Commercial Courts Act, 2015 (4 of 2016):"

    # Save document
    file_name = "Mediation_Application_Form.docx"
    doc.save(file_name)

    return send_file(file_name, as_attachment=True)


if __name__ == "__main__":
    app.run()
